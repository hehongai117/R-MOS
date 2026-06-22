"""Generate R-MOS product feature document in Word format."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Global style tweaks ──────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "微软雅黑"
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(4)

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.name = "微软雅黑"
    hs.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    hs.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    if level == 1:
        hs.font.size = Pt(22)
    elif level == 2:
        hs.font.size = Pt(16)
    else:
        hs.font.size = Pt(13)


def add_heading(text, level=1):
    doc.add_heading(text, level=level)


def add_para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p


def add_bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    run = p.add_run(text)
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.27 * level)
    return p


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = str(val)
    doc.add_paragraph()  # spacing


# =====================================================================
# COVER
# =====================================================================
for _ in range(6):
    doc.add_paragraph()

cover_title = doc.add_paragraph()
cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = cover_title.add_run("R-MOS 数字孪生维保智能体")
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

cover_sub = doc.add_paragraph()
cover_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = cover_sub.add_run("Robot Maintenance Operating System")
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()

cover_desc = doc.add_paragraph()
cover_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = cover_desc.add_run("产品功能介绍说明书")
run.font.size = Pt(20)

for _ in range(4):
    doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run("版本: v0.3.0  |  2026 年 5 月")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# =====================================================================
# 目录占位
# =====================================================================
add_heading("目 录", level=1)
add_para("（请在 Word 中插入自动目录：引用 → 目录 → 自动目录）")
doc.add_page_break()

# =====================================================================
# 1. 产品概述
# =====================================================================
add_heading("一、产品概述", level=1)

add_heading("1.1 产品定位", level=2)
add_para(
    "R-MOS（Robot Maintenance Operating System）是一套面向机器人维保培训与智能运维的"
    "全栈数字孪生平台。系统以高精度 3D 数字孪生为核心，融合多智能体 AI、实时遥测监控、"
    "标准化操作程序（SOP）执行引擎和知识库管理，为学员、教师和运维工程师提供沉浸式、"
    "可量化、可追溯的维保学习与实操环境。"
)

add_heading("1.2 核心价值", level=2)
add_bullet("沉浸式 3D 数字孪生 — 23 自由度仿人机器人 ATOM-01 高精度模型，支持爆炸图拆解、实时关节联动")
add_bullet("AI 驱动的智能诊断 — 规则引擎 + LLM 增强的故障诊断流水线，秒级定位故障根因")
add_bullet("裁决级 SOP 执行 — 基于约束图的操作裁决引擎，实时校验每一步操作的合规性")
add_bullet("五维技能画像 — 安全规范、步骤规范性、操作精度、时间效率、工具使用五个维度持续追踪")
add_bullet("全链路证据留存 — SHA-256 哈希封存的证据包，满足审计与合规要求")
add_bullet("多层 LLM 回退 — DeepSeek → MiniMax → Mock 三级回退链，保障 AI 能力永远可用")

add_heading("1.3 技术架构", level=2)
add_table(
    ["层次", "技术选型", "说明"],
    [
        ["前端", "React 18 + TypeScript + Vite", "单页应用，Zustand 状态管理"],
        ["3D 引擎", "React Three Fiber + Three.js", "364 个 GLB 模型资产"],
        ["后端", "FastAPI + SQLAlchemy 2.0 (Async)", "异步高性能 API 服务"],
        ["数据库", "PostgreSQL 14+ / SQLite (Dev)", "Alembic 迁移管理"],
        ["实时通信", "WebSocket", "5Hz 遥测推送"],
        ["AI/LLM", "DeepSeek / MiniMax / Mock", "统一路由 + 回退链"],
        ["适配器", "Mock / Gazebo / Real (ROS2)", "工厂模式可插拔"],
    ],
)

doc.add_page_break()

# =====================================================================
# 2. 3D 数字孪生系统
# =====================================================================
add_heading("二、3D 数字孪生系统", level=1)

add_heading("2.1 ATOM-01 机器人模型", level=2)
add_para(
    "系统内置 ATOM-01 仿人机器人高精度三维模型，具备 23 个自由度（DOF），"
    "包含 364 个独立零部件 GLB 模型资产。支持从整机总览（L0）到子系统拆解（L1）"
    "再到零件级细节（L2）的三级层次浏览。"
)

add_heading("2.2 爆炸图与拆解动画", level=2)
add_bullet("可配置爆炸系数（0~1 连续调节），平滑爆炸/收拢动画")
add_bullet("L1/L2 层级隔离聚焦，自动隐藏无关部件")
add_bullet("准 CAD 装配视图：默认视角、躯干维护视角、Left Knee Service 等预设")
add_bullet("分步爆炸步骤播放（如 Left Knee Service 共 7 步）")

add_heading("2.3 实时关节联动", level=2)
add_para(
    "3D 模型与后端 WebSocket 遥测数据实时联动（5Hz），关节状态实时驱动模型姿态。"
    "故障关节以红色闪烁标识，重点观测部位以绿色高亮。支持点击关节查看详细参数"
    "（位置、速度、扭矩、电流、温度）。"
)

add_heading("2.4 零件信息面板", level=2)
add_bullet("BOM 编码与物料清单")
add_bullet("所需工具列表")
add_bullet("技术规格参数")
add_bullet("螺丝信息（类型、数量、扭矩要求）")

doc.add_page_break()

# =====================================================================
# 3. 实时监控系统
# =====================================================================
add_heading("三、实时监控系统", level=1)

add_heading("3.1 监控仪表盘", level=2)
add_para("实时监控页以数字孪生为中心，联动展示机器人全方位状态：")
add_table(
    ["模块", "指标", "说明"],
    [
        ["机器人态势", "电池 / 核心温度 / 活跃故障 / 关节路数", "总体健康概览"],
        ["姿态与运动", "加速度 XYZ / 角速度 XYZ", "IMU 传感器数据"],
        ["电源与载荷", "主电压 / 逻辑电压 / 足底压力", "电压总线与压力传感器"],
        ["重点关节", "位置 / 速度 / 扭矩 / 电流 / 温度", "优先显示高温、异常关节"],
        ["故障定位", "故障码 + 关节名称", "联动 3D 模型高亮"],
    ],
)

add_heading("3.2 故障告警卡片", level=2)
add_para(
    "当关节温度超过 70°C 或主电压低于 20V 时，系统自动在监控页顶部弹出故障告警卡片"
    "（FaultAlertCard），显示故障类型、受影响关节、当前值与阈值。支持两个快捷操作："
)
add_bullet('"一键诊断" — 自动跳转至 AI 诊断工作台，携带故障上下文')
add_bullet('"忽略" — 暂时隐藏该告警')

add_heading("3.3 WebSocket 遥测", level=2)
add_bullet("5Hz 实时推送频率")
add_bullet("心跳检测与自动重连（指数退避，最多 10 次重试）")
add_bullet("数据新鲜度检测与过期警告")
add_bullet("连接状态指示器（已连接/断开/重连中）")

doc.add_page_break()

# =====================================================================
# 4. 故障诊断流水线
# =====================================================================
add_heading("四、AI 故障诊断流水线", level=1)

add_heading("4.1 诊断架构", level=2)
add_para("故障诊断采用「规则引擎 + LLM 增强」的两阶段流水线架构：")
add_bullet("第一阶段：规则引擎 — 基于阈值的确定性诊断，响应时间 < 10ms")
add_bullet("第二阶段：LLM 增强 — 大语言模型生成诊断推理报告（可选，最佳努力）")

add_heading("4.2 诊断阈值", level=2)
add_table(
    ["指标", "阈值", "对应故障"],
    [
        ["关节温度", "≥ 70°C", "E001_OVERHEAT（过热）"],
        ["主电压", "< 20V", "E003_VOLTAGE_DROP（电压跌落）"],
        ["位置误差", "≥ 0.10 rad", "E005_LOOSE（松动）"],
    ],
)

add_heading("4.3 三个标准故障场景", level=2)

add_para("E001_OVERHEAT — 关节过热（初级）", bold=True)
add_bullet("影响关节：腰部关节")
add_bullet("升温曲线：30 秒内温度上升 35°C，伴随扭矩噪声 ±0.5 Nm")
add_bullet("推荐 SOP：关节过热应急处理（4 步骤）")

add_para("E005_LOOSE — 关节松动（中级）", bold=True)
add_bullet("影响关节：肘部关节")
add_bullet("症状：位置误差 +0.14 rad，振动 +2.0")
add_bullet("推荐 SOP：关节松动检修（6 步骤）")

add_para("E003_VOLTAGE_DROP — 电压跌落复合故障（高级）", bold=True)
add_bullet("影响关节：肩部、肘部")
add_bullet("症状：主电压 -5V，引发二次过热 +25°C")
add_bullet("复合触发：自动级联触发 E001_OVERHEAT")
add_bullet("推荐 SOP：电压跌落复合故障处理（8 步骤）")

add_heading("4.4 诊断流水线 API", level=2)
add_table(
    ["端点", "方法", "功能"],
    [
        ["/pipeline/diagnose", "POST", "遥测数据诊断"],
        ["/pipeline/tasks/from-diagnosis", "POST", "从诊断结果创建维保任务"],
        ["/pipeline/executions/{id}/steps/complete", "POST", "步骤完成上报"],
        ["/pipeline/executions/{id}/complete", "POST", "任务完成与报告触发"],
    ],
)

doc.add_page_break()

# =====================================================================
# 5. SOP 维保执行系统
# =====================================================================
add_heading("五、SOP 维保执行系统", level=1)

add_heading("5.1 裁决级 SOP 播放器", level=2)
add_para(
    "系统内置裁决级 SOP 播放器（SOPPlayerAdjudicated），基于约束图的操作裁决引擎，"
    "实时校验每一步操作的合规性。支持教学模式、考试模式和维保模式三种执行模式。"
)

add_heading("5.2 约束系统", level=2)
add_para("七种零件约束关系确保拆装顺序的物理正确性：")
add_table(
    ["约束类型", "含义", "示例"],
    [
        ["FASTENED_BY", "螺丝紧固", "面板由 4 颗 M3 螺丝固定"],
        ["COVERED_BY", "覆盖遮挡", "主板被上盖板覆盖"],
        ["BLOCKED_BY", "几何干涉", "电机被支架阻挡"],
        ["LOCKED_BY", "机械锁定", "轴承由卡簧锁定"],
        ["HINGED_TO", "铰链连接", "可旋转但不可拆卸"],
        ["WIRED_TO", "线缆连接", "电机接线需先断开"],
        ["PLUGGED_TO", "插接连接", "信号线插头"],
    ],
)

add_heading("5.3 裁决结果", level=2)
add_table(
    ["结果", "含义", "处理"],
    [
        ["ALLOWED", "操作允许", "执行并推进到下一步"],
        ["BLOCKED", "硬约束违反", "阻止操作，提示原因"],
        ["WARNING", "软约束提醒", "允许继续但给出警告"],
        ["TOOL_MISMATCH", "工具不匹配", "提示更换正确工具"],
        ["INCOMPLETE", "操作未完成", "需完成当前操作"],
    ],
)

add_heading("5.4 SOP 执行状态机", level=2)
add_para("IDLE → PRECONDITION_CHECK → EXECUTING → VALIDATION → COMPLETE")
add_para("任何阶段均可进入 BLOCKED 状态，BLOCKED 后可重试或回退。")

add_heading("5.5 Pipeline 同步", level=2)
add_para(
    "SOP 步骤完成时自动向后端 Pipeline API 同步步骤结果（含证据类型、耗时），"
    "任务全部完成后自动触发任务完成与报告生成，并跳转至维保报告页。"
)

doc.add_page_break()

# =====================================================================
# 6. 多智能体系统
# =====================================================================
add_heading("六、多智能体 AI 系统", level=1)

add_heading("6.1 Agent 工作台", level=2)
add_para(
    "AI 诊断工作台提供维保编排、知识查询与轨迹追踪统一入口。支持自然语言交互，"
    "可通过快捷按钮（通用问答、派单维保、诊断问题、知识查询、知识记录、训练指导）"
    "快速发起意图。"
)

add_heading("6.2 多智能体架构", level=2)
add_table(
    ["智能体", "职责", "能力"],
    [
        ["Orchestrator", "协调器", "任务分发与结果汇聚"],
        ["Diagnoser", "诊断专家", "根因分析、干预建议"],
        ["Coach", "教练", "实时引导与反馈"],
        ["Curriculum", "课程规划", "学习路径管理"],
    ],
)

add_heading("6.3 诊断专家能力", level=2)
add_para("诊断专家（Diagnoser）能识别以下根因类型：")
add_bullet("概念误解（Concept Misunderstanding）")
add_bullet("习惯问题（Habit Issue）")
add_bullet("注意力缺失（Attention Deficit）")
add_bullet("工具选择错误（Tool Selection Error）")
add_bullet("步骤顺序错误（Sequence Error）")

add_para("并推荐对应干预措施：讲解（Explain）、练习（Practice）、演示（Demo）、检查点（Checkpoint）。")

add_heading("6.4 LLM 服务", level=2)
add_para("统一 LLM 路由器支持多个大语言模型提供商：")
add_table(
    ["提供商", "用途", "接入方式"],
    [
        ["DeepSeek", "主力推理", "OpenAI 兼容 SDK"],
        ["MiniMax", "回退备选", "HTTP REST API"],
        ["Mock", "开发测试", "确定性响应，无外部依赖"],
    ],
)
add_para("三级回退链确保 AI 能力始终可用：DeepSeek → MiniMax → Mock。")
add_para("LLM 健康检查端点：GET /api/v1/llm/health，实时监控各 Provider 状态。")

doc.add_page_break()

# =====================================================================
# 7. 训练与评估系统
# =====================================================================
add_heading("七、训练与评估系统", level=1)

add_heading("7.1 训练项目生成", level=2)
add_para(
    "系统支持 AI 驱动的训练项目生成（ProjectGenerator），通过双路径检索"
    "（知识库 + 个人记忆）为学员定制训练方案。支持 SSE 流式响应，实时展示生成过程。"
    "生成内容包括：步骤清单、工具检查表、裁决配置、机器人规格等。"
)

add_heading("7.2 训练会话管理", level=2)
add_para("训练会话状态机：")
add_bullet("active → paused → active（暂停/恢复）")
add_bullet("active → submitted（手动/超时/教师提交）")
add_bullet("active → abandoned（学员放弃）")
add_bullet("active → expired（48 小时超时）")

add_heading("7.3 五维技能画像", level=2)
add_para("系统持续追踪学员的五个维度技能分数（0~100 分）：")
add_table(
    ["维度", "英文", "评估内容"],
    [
        ["安全规范执行", "Safety", "是否遵守安全隔离、断电等规范"],
        ["步骤规范性", "Procedure", "操作步骤顺序与完整性"],
        ["操作精度", "Precision", "扭矩、位置、力度等精度"],
        ["时间效率", "Efficiency", "操作耗时与标准时间对比"],
        ["工具使用规范", "Tools", "工具选择与使用正确性"],
    ],
)
add_para("技能等级从 L1 到 L3 逐步进阶，系统自动评估认证资格。")

add_heading("7.4 评分引擎", level=2)
add_para("MVP 评分规则（满分 100 分）：")
add_bullet("跳过步骤：每次 -5 分")
add_bullet("错误操作：每次 -10 分")
add_bullet("超时：-15 分")
add_bullet("异常快照：每个 -15 分")

add_para("分数按四个维度加权：")
add_table(
    ["维度", "权重", "说明"],
    [
        ["执行正确性", "40%", "操作是否符合 SOP 要求"],
        ["安全合规性", "30%", "是否遵守安全规范"],
        ["程序规范性", "20%", "步骤顺序与完整度"],
        ["时间效率", "10%", "完成速度"],
    ],
)

add_heading("7.5 薄弱环节追踪", level=2)
add_para(
    "系统自动识别学员薄弱步骤（Weak Steps），记录失败模式与频次。"
    "结合训练记忆（Training Memory）采用滑动平均更新技能分数，"
    "并预计算下次推荐训练内容。"
)

doc.add_page_break()

# =====================================================================
# 8. 知识管理系统
# =====================================================================
add_heading("八、知识管理系统", level=1)

add_heading("8.1 知识文档管理", level=2)
add_para(
    "系统内置结构化知识库，支持故障维修手册、操作规范、参数手册等多类文档。"
    "每篇文档携带故障标签（fault_tags）和 SOP 标签（sop_tags），支持通配符匹配。"
)

add_heading("8.2 混合检索策略", level=2)
add_bullet("标签精确匹配 — 通过故障类型标签快速定位相关文档（高精度）")
add_bullet("文本搜索 — 关键词模糊匹配（高召回）")
add_bullet("向量语义搜索 — pgvector 嵌入检索（预留接口）")

add_heading("8.3 知识治理", level=2)
add_para("知识文档生命周期管理：")
add_bullet("状态流转：DRAFT → PENDING → APPROVED → EXPIRED")
add_bullet("风险等级分类（R0~R3）")
add_bullet("适用范围定义（设备型号、零件类型、版本范围）")
add_bullet("禁忌项追踪")
add_bullet("过期管理（时间/使用量/条件触发）")
add_bullet("置信度指标（证据数量、成功率、审核员评分）")

add_heading("8.4 种子知识库", level=2)
add_para("系统预置 5 篇知识文档：")
add_table(
    ["文档", "类型", "标签"],
    [
        ["关节过热维修手册", "故障专用", "E001_OVERHEAT"],
        ["关节松动维修手册", "故障专用", "E005_LOOSE"],
        ["电压系统维修手册", "故障专用", "E003_VOLTAGE_DROP"],
        ["安全操作通用规范", "通用指南", "* (全局)"],
        ["ATOM-01 结构参数手册", "通用指南", "* (全局)"],
    ],
)

doc.add_page_break()

# =====================================================================
# 9. 教学管理系统
# =====================================================================
add_heading("九、教学管理系统", level=1)

add_heading("9.1 教师功能", level=2)
add_bullet("创建和管理班级、课程、学员名册")
add_bullet("布置维保作业（Assignment），指定 SOP 和评分策略")
add_bullet("实时监控学员操作进度")
add_bullet("查看学员技能画像和薄弱环节")
add_bullet("审批队列管理")

add_heading("9.2 教学策略配置", level=2)
add_para("教师可通过引导策略（GuidancePolicy）精细控制教学行为：")
add_table(
    ["配置项", "说明", "默认值"],
    [
        ["ghost_hand", "AI 辅助手（自动引导）", "开启"],
        ["hint_button", "提示按钮可用性", "开启"],
        ["error_detail", "错误详情可见性", "显示"],
        ["max_retry", "最大重试次数", "3"],
        ["blind_steps", "盲步（考试模式隐藏步骤）", "关闭"],
        ["competition_mode", "竞赛模式", "关闭"],
    ],
)

add_heading("9.3 作业与考试", level=2)
add_para(
    "支持多次尝试（可配置最大尝试次数），尝试状态流转："
    "in_progress → completed → graded | abandoned。"
    "考试模式支持倒计时（默认 60 分钟）和盲步功能。"
)

doc.add_page_break()

# =====================================================================
# 10. 证据与审计系统
# =====================================================================
add_heading("十、证据与审计系统", level=1)

add_heading("10.1 证据包", level=2)
add_para(
    "系统在任务完成时自动生成证据包（Evidence Bundle），采用 SHA-256 哈希封存，"
    "确保数据完整性不可篡改。证据包类型包括：任务执行、评估记录、教学尝试。"
)

add_heading("10.2 证据项", level=2)
add_bullet("内容 URI 与哈希值")
add_bullet("MIME 类型追踪")
add_bullet("观察时间与采集时间")
add_bullet("机器标签（自动分类）")

add_heading("10.3 审计日志", level=2)
add_para(
    "全链路审计追踪：操作日志、评估变更、审批流程、"
    "诊断轨迹均有完整记录，支持按时间、用户、操作类型检索。"
)

doc.add_page_break()

# =====================================================================
# 11. 仿真与适配器
# =====================================================================
add_heading("十一、仿真与适配器系统", level=1)

add_heading("11.1 适配器模式", level=2)
add_para(
    "系统采用工厂模式的可插拔适配器架构，支持三种机器人连接方式："
)
add_table(
    ["适配器", "状态", "说明"],
    [
        ["MockRobotAdapter", "已实现", "合成遥测数据，支持故障注入"],
        ["GazeboAdapter", "规划中", "Gazebo 物理仿真器集成"],
        ["RealAdapter", "规划中", "ROS2 实机连接"],
    ],
)

add_heading("11.2 故障注入", level=2)
add_para("Mock 适配器支持 5 种故障注入：")
add_table(
    ["故障码", "名称", "效果"],
    [
        ["E001_OVERHEAT", "过热", "温度 +30°C，扭矩 -30%"],
        ["E002_STALL", "卡死", "速度归零，位置冻结"],
        ["E003_VOLTAGE_DROP", "电压跌落", "电池 -50%，扭矩 -50%"],
        ["E004_SENSOR_FAILURE", "传感器故障", "数据加入噪声"],
        ["E005_JOINT_LOOSE", "关节松动", "位置噪声，扭矩 -70%"],
    ],
)

doc.add_page_break()

# =====================================================================
# 12. 系统页面总览
# =====================================================================
add_heading("十二、系统页面总览", level=1)

add_table(
    ["页面", "路由", "角色", "核心功能"],
    [
        ["实时监控", "/monitor", "全员", "3D 数字孪生 + 遥测仪表盘 + 故障告警"],
        ["AI 诊断工作台", "/agent/workbench", "全员", "多智能体交互 + 诊断面板"],
        ["维保练习工作台", "/maintenance", "学员", "SOP 执行 + 3D 爆炸图引导"],
        ["我的任务", "/my-tasks", "学员", "任务列表（待完成/进行中/已完成）"],
        ["自主练习", "/scenarios", "学员", "场景选择（入门/进阶/高级）"],
        ["维保报告", "/reports", "全员", "任务报告 + 评分明细"],
        ["我的技能", "/student/skills", "学员", "五维雷达图 + 技能等级"],
        ["3D 展示", "/atom01", "全员", "ATOM-01 模型展示 + 动画"],
        ["知识库", "/knowledge", "全员", "知识文档浏览与搜索"],
        ["教学管理", "/teaching/*", "教师", "班级/作业/学员管理"],
        ["管理控制台", "/admin/console", "管理员", "系统健康 + 用户管理 + 审批"],
    ],
)

doc.add_page_break()

# =====================================================================
# 13. API 接口
# =====================================================================
add_heading("十三、主要 API 接口", level=1)

add_heading("13.1 核心接口", level=2)
add_table(
    ["接口", "方法", "功能"],
    [
        ["/api/v1/health", "GET", "系统健康检查"],
        ["/api/v1/sops", "GET/POST", "SOP 列表与创建"],
        ["/api/v1/tasks", "POST", "创建维保任务"],
        ["/api/v1/tasks/{id}/start", "POST", "启动任务"],
        ["/api/v1/tasks/{id}/step", "POST", "执行步骤"],
        ["/api/v1/tasks/{id}/report", "GET", "获取任务报告"],
        ["/ws/robot/status", "WS", "实时遥测推送（5Hz）"],
    ],
)

add_heading("13.2 Pipeline 接口", level=2)
add_table(
    ["接口", "方法", "功能"],
    [
        ["/api/v1/pipeline/diagnose", "POST", "故障诊断"],
        ["/api/v1/pipeline/tasks/from-diagnosis", "POST", "诊断 → 任务创建"],
        ["/api/v1/pipeline/executions/{id}/steps/complete", "POST", "步骤完成"],
        ["/api/v1/pipeline/executions/{id}/complete", "POST", "任务完成"],
        ["/api/v1/llm/health", "GET", "LLM 健康检查"],
    ],
)

add_heading("13.3 教学接口", level=2)
add_table(
    ["接口", "方法", "功能"],
    [
        ["/api/v1/training/sessions", "POST", "创建训练会话"],
        ["/api/v1/training/sessions/{id}/submit", "POST", "提交训练"],
        ["/api/v1/training/feedback/{id}", "GET", "获取 AI 反馈"],
        ["/api/v1/students/{id}/profile", "GET", "获取技能画像"],
    ],
)

doc.add_page_break()

# =====================================================================
# 14. 角色与权限
# =====================================================================
add_heading("十四、角色与权限", level=1)

add_table(
    ["角色", "权限范围", "典型操作"],
    [
        ["管理员 (Admin)", "全局管理权限", "用户管理、系统配置、审批、知识治理"],
        ["教师 (Teacher)", "教学域管理权限", "创建作业、监控学员、评分、审批高风险操作"],
        ["学员 (Student)", "学习域操作权限", "执行任务、查看报告、自主练习、知识查阅"],
    ],
)

add_para(
    "系统采用 JWT 认证 + RBAC 授权机制，支持路由级保护（ProtectedRoute）"
    "和资源级权限控制（authz_guard）。"
)

doc.add_page_break()

# =====================================================================
# 15. 部署
# =====================================================================
add_heading("十五、部署与运行", level=1)

add_heading("15.1 后端启动", level=2)
add_para("cd r-mos-backend && source .venv/bin/activate")
add_para("pip install -r requirements.txt")
add_para("alembic upgrade head")
add_para("python main.py")

add_heading("15.2 前端启动", level=2)
add_para("cd r-mos-frontend && npm install && npm run dev")

add_heading("15.3 环境变量", level=2)
add_table(
    ["变量", "说明", "默认值"],
    [
        ["DATABASE_URL", "数据库连接", "postgresql+asyncpg://..."],
        ["ROBOT_ADAPTER_TYPE", "适配器类型", "mock"],
        ["DEEPSEEK_API_KEY", "DeepSeek API 密钥", "(空)"],
        ["MINIMAX_API_KEY", "MiniMax API 密钥", "(空)"],
        ["LLM_PRIMARY_PROVIDER", "主力 LLM 提供商", "deepseek"],
        ["LLM_FALLBACK_PROVIDER", "回退 LLM 提供商", "minimax"],
        ["LLM_ENABLE_MOCK_FALLBACK", "启用 Mock 兜底", "True"],
    ],
)

# =====================================================================
# Save
# =====================================================================
out_path = os.path.join(os.path.dirname(__file__), "..", "R-MOS 数字孪生维保智能体功能介绍.docx")
out_path = os.path.abspath(out_path)
doc.save(out_path)
print(f"Document saved to: {out_path}")
